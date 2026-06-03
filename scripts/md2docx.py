#!/usr/bin/env python3
"""Convert paper.md to paper.docx with proper Word equations (OOMML)."""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

INPUT = "/root/dslm/deepsleep/docs/paper.md"
OUTPUT = "/root/dslm/deepsleep/docs/paper.docx"

# ═══════════════════════════════════════════════════════════════
# LaTeX → OOMML (Office Math Markup Language) converter
# ═══════════════════════════════════════════════════════════════

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _m(tag):
    """Create a math namespace tag."""
    return f"{{{MATH_NS}}}{tag}"


def _w(tag):
    """Create a word namespace tag."""
    return f"{{{W_NS}}}{tag}"


def latex_to_oomml(latex_str):
    """Convert a LaTeX math string to OOMML XML element.

    Handles: Greek letters, subscripts, superscripts, fractions, sums,
    products, integrals, square roots, hats, bars, text, common symbols.
    """
    s = latex_str.strip()
    s = re.sub(r'\\tag\{(\d+\w?)\}', '', s)  # remove \tag{N}
    s = re.sub(r'\\label\{[^}]*\}', '', s)
    s = re.sub(r'\\quad\s*', '  ', s)
    s = s.strip()

    try:
        oMath = etree.Element(_m("oMath"))
        _parse_expr(s, oMath, 0, len(s))[0]
        return oMath
    except Exception:
        # Fallback: render as italic text
        return None


def _parse_expr(s, parent, pos, end):
    """Parse expression and append OOMML nodes to parent. Returns (new_pos)."""
    nodes_added = 0
    while pos < end:
        # Skip whitespace
        if s[pos] == ' ':
            pos += 1
            continue

        # Backslash commands
        if s[pos] == '\\':
            pos = _parse_command(s, parent, pos, end)
            nodes_added += 1
            continue

        # Superscript ^
        if s[pos] == '^':
            pos += 1
            if nodes_added == 0:
                r = etree.SubElement(parent, _m("r"))
                etree.SubElement(r, _m("t")).text = "^"
                nodes_added += 1
                continue
            # Wrap last child in sSup
            last = parent[-1]
            parent.remove(last)
            sSup = etree.SubElement(parent, _m("sSup"))
            sSupPr = etree.SubElement(sSup, _m("sSupPr"))
            etree.SubElement(sSup, _m("e")).append(last)
            sup = etree.SubElement(sSup, _m("sup"))
            pos = _parse_arg(s, sup, pos, end)
            continue

        # Subscript _
        if s[pos] == '_':
            pos += 1
            if nodes_added == 0:
                r = etree.SubElement(parent, _m("r"))
                etree.SubElement(r, _m("t")).text = "_"
                nodes_added += 1
                continue
            last = parent[-1]
            parent.remove(last)
            sSub = etree.SubElement(parent, _m("sSub"))
            etree.SubElement(sSub, _m("sSubPr"))
            etree.SubElement(sSub, _m("e")).append(last)
            sub = etree.SubElement(sSub, _m("sub"))
            pos = _parse_arg(s, sub, pos, end)
            continue

        # Curly braces - just skip them as group delimiters
        if s[pos] == '{':
            depth = 1
            start = pos + 1
            pos += 1
            while pos < end and depth > 0:
                if s[pos] == '{': depth += 1
                elif s[pos] == '}': depth -= 1
                pos += 1
            group_content = s[start:pos-1]
            pos2 = _parse_expr(group_content, parent, 0, len(group_content))[0]
            continue

        if s[pos] == '}':
            return pos + 1, nodes_added

        # Regular character
        r = etree.SubElement(parent, _m("r"))
        rPr = etree.SubElement(r, _m("rPr"))
        sty = etree.SubElement(rPr, _m("sty"))
        sty.set(_w("val"), "i")  # italic
        etree.SubElement(r, _m("t")).text = s[pos]
        pos += 1
        nodes_added += 1

    return pos, nodes_added


def _parse_arg(s, parent, pos, end):
    """Parse a single argument (char, {group}, or command)."""
    if pos >= end:
        return pos
    if s[pos] == '{':
        depth = 1
        pos += 1
        start = pos
        while pos < end and depth > 0:
            if s[pos] == '{': depth += 1
            elif s[pos] == '}': depth -= 1
            pos += 1
        group = s[start:pos-1]
        _parse_expr(group, parent, 0, len(group))
        return pos
    if s[pos] == '\\':
        return _parse_command(s, parent, pos, end)
    # Single char
    r = etree.SubElement(parent, _m("r"))
    rPr = etree.SubElement(r, _m("rPr"))
    sty = etree.SubElement(rPr, _m("sty"))
    sty.set(_w("val"), "i")
    etree.SubElement(r, _m("t")).text = s[pos]
    return pos + 1


def _parse_command(s, parent, pos, end):
    """Parse a LaTeX command starting with backslash."""
    pos += 1  # skip backslash
    if pos >= end:
        return pos

    # Single-char commands
    if not s[pos].isalpha():
        ch = s[pos]
        if ch in ('{', '}', '^', '_', ' ', '#', '%', '&', '|'):
            r = etree.SubElement(parent, _m("r"))
            etree.SubElement(r, _m("t")).text = ch
            return pos + 1
        r = etree.SubElement(parent, _m("r"))
        etree.SubElement(r, _m("t")).text = '\\' + ch
        return pos + 1

    # Read command name
    start = pos
    while pos < end and s[pos].isalpha():
        pos += 1
    cmd = s[start:pos]

    # Greek letters
    GREEK = {
        'alpha': 'α', 'beta': 'β', 'gamma': 'γ', 'delta': 'δ', 'epsilon': 'ε',
        'zeta': 'ζ', 'eta': 'η', 'theta': 'θ', 'iota': 'ι', 'kappa': 'κ',
        'lambda': 'λ', 'mu': 'μ', 'nu': 'ν', 'xi': 'ξ', 'pi': 'π', 'rho': 'ρ',
        'sigma': 'σ', 'tau': 'τ', 'upsilon': 'υ', 'phi': 'φ', 'chi': 'χ',
        'psi': 'ψ', 'omega': 'ω', 'Gamma': 'Γ', 'Delta': 'Δ', 'Theta': 'Θ',
        'Lambda': 'Λ', 'Xi': 'Ξ', 'Pi': 'Π', 'Sigma': 'Σ', 'Phi': 'Φ',
        'Psi': 'Ψ', 'Omega': 'Ω', 'varepsilon': 'ε', 'varphi': 'φ',
        'varkappa': 'κ', 'ell': 'ℓ',
    }
    if cmd in GREEK:
        r = etree.SubElement(parent, _m("r"))
        etree.SubElement(r, _m("t")).text = GREEK[cmd]
        return pos

    # Symbols
    SYMBOLS = {
        'infty': '∞', 'partial': '∂', 'nabla': '∇', 'cdot': '·', 'times': '×',
        'pm': '±', 'mp': '∓', 'div': '÷', 'approx': '≈', 'equiv': '≡',
        'neq': '≠', 'leq': '≤', 'geq': '≥', 'll': '≪', 'gg': '≫',
        'sim': '~', 'simeq': '≃', 'propto': '∝', 'parallel': '∥',
        'perp': '⊥', 'angle': '∠', 'circ': '∘', 'bullet': '•',
        'oplus': '⊕', 'otimes': '⊗', 'star': '★', 'dag': '†',
        'hbar': 'ℏ', 'forall': '∀', 'exists': '∃', 'nabla': '∇',
        'to': '→', 'gets': '←', 'leftrightarrow': '↔', 'uparrow': '↑',
        'downarrow': '↓', 'Rightarrow': '⇒', 'Leftarrow': '⇐',
        'mapsto': '↦', 'implies': '⟹', 'iff': '⟺',
        'ldots': '…', 'cdots': '⋯', 'vdots': '⋮', 'ddots': '⋱',
        'sum': '∑', 'prod': '∏', 'coprod': '∐', 'int': '∫',
        'iint': '∬', 'iiint': '∭', 'oint': '∮',
        'bigcup': '⋃', 'bigcap': '⋂', 'bigsqcup': '⊔',
        'bigoplus': '⊕', 'bigotimes': '⊗', 'bigodot': '⊙',
        'sqrt': None, 'frac': None, 'hat': None, 'bar': None,
        'tilde': None, 'vec': None, 'dot': None, 'ddot': None,
        'overline': None, 'underline': None, 'overbrace': None,
        'text': None, 'mathrm': None, 'mathbf': None, 'mathcal': None,
        'mathbb': None, 'mathfrak': None, 'left': None, 'right': None,
        'displaystyle': None, 'textstyle': None, 'limits': None,
        'operatorname': None, 'log': 'log', 'ln': 'ln', 'exp': 'exp',
        'sin': 'sin', 'cos': 'cos', 'tan': 'tan', 'max': 'max', 'min': 'min',
        'arg': 'arg', 'dim': 'dim', 'det': 'det', 'ker': 'ker',
        'lim': 'lim', 'sup': 'sup', 'inf': 'inf', 'limsup': 'lim sup',
        'liminf': 'lim inf', 'Pr': 'Pr', 'Var': 'Var', 'Cov': 'Cov',
        'E': 'E', 'P': 'P',
    }

    if cmd in SYMBOLS:
        val = SYMBOLS[cmd]
        if val is None:
            if cmd == 'sqrt':
                pos = _parse_arg(s, parent, pos, end)
                # wrap last child in rad
                last = parent[-1]
                parent.remove(last)
                rad = etree.SubElement(parent, _m("rad"))
                etree.SubElement(rad, _m("radPr"))
                etree.SubElement(rad, _m("deg"))
                etree.SubElement(rad, _m("e")).append(last)
                return pos
            elif cmd == 'frac':
                num = etree.SubElement(parent, _m("f"))
                etree.SubElement(num, _m("fPr"))
                numElem = etree.SubElement(num, _m("num"))
                pos = _parse_arg(s, numElem, pos, end)
                denElem = etree.SubElement(num, _m("den"))
                pos = _parse_arg(s, denElem, pos, end)
                return pos
            elif cmd in ('hat', 'bar', 'tilde', 'vec', 'dot', 'ddot', 'overline', 'underline'):
                ACCENT_MAP = {
                    'hat': '̂', 'bar': '̄', 'tilde': '̃', 'vec': '→',
                    'dot': '̇', 'ddot': '̈', 'overline': '̄', 'underline': '̲',
                }
                last_before = len(parent)
                pos = _parse_arg(s, parent, pos, end)
                # Wrap newly added element in accent
                if len(parent) > last_before:
                    accent_elem = parent[-1]
                    parent.remove(accent_elem)
                    acc = etree.SubElement(parent, _m("acc"))
                    etree.SubElement(acc, _m("accPr"))
                    etree.SubElement(acc, _m("e")).append(accent_elem)
                return pos
            elif cmd in ('text', 'mathrm', 'operatorname'):
                pos = _parse_arg(s, parent, pos, end)
                # Make last child non-italic
                if len(parent) > 0 and parent[-1].tag == _m("r"):
                    r = parent[-1]
                    for t in r.iter(_m("t")):
                        # remove italic
                        for rPr in r.findall(_m("rPr")):
                            r.remove(rPr)
                return pos
            elif cmd in ('mathbf', 'mathcal', 'mathbb', 'mathfrak'):
                pos = _parse_arg(s, parent, pos, end)
                return pos
            elif cmd in ('left', 'right', 'displaystyle', 'textstyle', 'limits'):
                return pos  # skip formatting commands
            return pos
        else:
            # Function name (log, sin, etc.)
            r = etree.SubElement(parent, _m("r"))
            # Remove italic for function names
            etree.SubElement(r, _m("t")).text = val
            return pos

    # Unknown command - render as text
    r = etree.SubElement(parent, _m("r"))
    etree.SubElement(r, _m("t")).text = '\\' + cmd
    return pos


def add_math_paragraph(doc, latex_str, center=True):
    """Add a paragraph containing a Word equation."""
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    oMath = latex_to_oomml(latex_str)
    if oMath is not None:
        # Wrap in oMathPara for display math
        oMathPara = etree.SubElement(p._element, _m("oMathPara"))
        oMathPara.append(oMath)
    else:
        # Fallback: plain italic text
        run = p.add_run(latex_str)
        run.italic = True
        run.font.name = 'Cambria Math'
        run.font.size = Pt(11)
    return p


def add_inline_math(run_element, latex_str):
    """Add inline math to an existing run's paragraph."""
    p = run_element.getparent().getparent()  # run -> r -> p
    oMath = latex_to_oomml(latex_str)
    if oMath is not None:
        # Insert math after the run
        r_elem = run_element.getparent()  # <w:r>
        r_elem.addnext(oMath)
    return None


# ═══════════════════════════════════════════════════════════════
# Markdown parser
# ═══════════════════════════════════════════════════════════════

def parse_md(filepath):
    """Parse markdown file into structured blocks."""
    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    blocks = []
    i = 0
    while i < len(lines):
        line = lines[i].rstrip('\n')

        # Display math $$...$$
        if line.strip().startswith('$$'):
            math_lines = [line.strip()[2:]]
            i += 1
            while i < len(lines):
                mline = lines[i].rstrip('\n')
                if mline.strip().endswith('$$'):
                    remaining = mline.strip()[:-2]
                    if remaining:
                        math_lines.append(remaining)
                    i += 1
                    break
                math_lines.append(mline.strip())
                i += 1
            blocks.append(('display_math', '\n'.join(math_lines)))
            continue

        # Headings
        m = re.match(r'^(#{1,4})\s+(.*)', line)
        if m:
            level = len(m.group(1))
            text = m.group(2)
            blocks.append(('heading', level, text))
            i += 1
            continue

        # Table
        if '|' in line and i + 1 < len(lines) and '---' in lines[i + 1]:
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                table_lines.append(lines[i].rstrip('\n'))
                i += 1
            blocks.append(('table', table_lines))
            continue

        # Horizontal rule
        if line.strip() == '---':
            i += 1
            continue

        # Empty line
        if line.strip() == '':
            i += 1
            continue

        # Figure placeholder: [图: filename]
        fm = re.match(r'^\[图:\s*(.+?)\]', line.strip())
        if fm:
            blocks.append(('figure', fm.group(1).strip()))
            i += 1
            continue

        # Block quote
        if line.startswith('>'):
            quote_text = line.lstrip('> ').strip()
            i += 1
            while i < len(lines) and lines[i].startswith('>'):
                quote_text += ' ' + lines[i].lstrip('> ').strip()
                i += 1
            blocks.append(('quote', quote_text))
            continue

        # Regular paragraph
        para_lines = [line]
        i += 1
        while i < len(lines):
            pl = lines[i].rstrip('\n')
            if pl.strip() == '' or pl.startswith('#') or pl.strip() == '---' or \
               pl.strip().startswith('$$') or (pl.startswith('|') and i + 1 < len(lines) and '---' in lines[i+1]) or \
               pl.strip().startswith('[图:'):
                break
            if pl.startswith('>'):
                break
            para_lines.append(pl)
            i += 1
        blocks.append(('paragraph', '\n'.join(para_lines)))

    return blocks


def add_formatted_text(paragraph, text):
    """Add text with inline formatting: **bold**, *italic*, `code`, $math$, [links]."""
    # Pattern: $...$ (inline math), **...** (bold), *...* (italic), `...` (code), [...](...) (link)
    pattern = r'(\$[^$]+\$|\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`|\[([^\]]+)\]\([^)]+\))'
    parts = re.split(pattern, text)

    for part in parts:
        if not part:
            continue
        if part.startswith('$') and part.endswith('$'):
            # Inline math
            latex = part[1:-1]
            oMath = latex_to_oomml(latex)
            if oMath is not None:
                paragraph._element.append(oMath)
            else:
                run = paragraph.add_run(latex)
                run.italic = True
                run.font.name = 'Cambria Math'
        elif part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
            run.font.name = '宋体'
            run.font.size = Pt(12)
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
            run.font.name = '宋体'
            run.font.size = Pt(12)
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(10)
        elif re.match(r'^\[([^\]]+)\]\([^)]+\)$', part):
            m = re.match(r'^\[([^\]]+)\]\(([^)]+)\)$', part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.name = '宋体'
                run.font.size = Pt(12)
        else:
            run = paragraph.add_run(part)
            run.font.name = '宋体'
            run.font.size = Pt(12)


# ═══════════════════════════════════════════════════════════════
# Main conversion
# ═══════════════════════════════════════════════════════════════

def main():
    print(f"Reading {INPUT}...")
    blocks = parse_md(INPUT)
    print(f"Parsed {len(blocks)} blocks.")

    doc = Document()

    # Set default style
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # Set East Asian font
    rPr = style.element.find(qn('w:rPr'))
    if rPr is None:
        rPr = style.element.makeelement(qn('w:rPr'), {})
        style.element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), '宋体')

    for block in blocks:
        btype = block[0]

        if btype == 'heading':
            level, text = block[1], block[2]
            # Remove markdown formatting from heading
            clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
            h = doc.add_heading(level=min(level, 4))
            run = h.add_run(clean)
            font_sizes = {1: 16, 2: 14, 3: 13, 4: 12}
            run.font.size = Pt(font_sizes.get(level, 12))
            run.font.name = '黑体'
            # Set East Asian font
            r = run._element
            rPr2 = r.find(qn('w:rPr'))
            if rPr2 is None:
                rPr2 = r.makeelement(qn('w:rPr'), {})
                r.insert(0, rPr2)
            rF = rPr2.find(qn('w:rFonts'))
            if rF is None:
                rF = rPr2.makeelement(qn('w:rFonts'), {})
                rPr2.append(rF)
            rF.set(qn('w:eastAsia'), '黑体')

        elif btype == 'display_math':
            add_math_paragraph(doc, block[1])

        elif btype == 'paragraph':
            text = block[1]
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(0.74)
            add_formatted_text(p, text)

        elif btype == 'figure':
            filename = block[1]
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(12)
            run = p.add_run(f'[图片占位: {filename}]')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(100, 100, 100)
            run.font.name = '宋体'
            # Add caption line
            # (caption is usually next paragraph block starting with *)

        elif btype == 'table':
            table_lines = block[1]
            # Parse table
            rows_data = []
            for tl in table_lines:
                if '---' in tl:
                    continue
                cells = [c.strip() for c in tl.strip('|').split('|')]
                rows_data.append(cells)

            if not rows_data:
                continue

            ncols = max(len(r) for r in rows_data)
            nrows = len(rows_data)

            table = doc.add_table(rows=nrows, cols=ncols)
            table.style = 'Table Grid'

            for ri, row in enumerate(rows_data):
                for ci, cell_text in enumerate(row):
                    if ci >= ncols:
                        break
                    cell = table.rows[ri].cells[ci]
                    cell.text = ''
                    cp = cell.paragraphs[0]
                    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    # Parse inline formatting in cells too
                    clean = re.sub(r'\$([^$]+)\$', r'\1', cell_text)  # simplify math in tables
                    clean = re.sub(r'\*\*([^*]+)\*\*', r'\1', clean)
                    clean = re.sub(r'\*([^*]+)\*', r'\1', clean)
                    run = cp.add_run(clean)
                    run.font.size = Pt(10)
                    run.font.name = '宋体'
                    if ri == 0:
                        run.bold = True

        elif btype == 'quote':
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.0)
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', block[1])
            text = re.sub(r'\*([^*]+)\*', r'\1', text)
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(80, 80, 80)
            run.font.name = '宋体'

    doc.save(OUTPUT)
    print(f"\nDone! Saved to {OUTPUT}")
    print(f"File size: {os.path.getsize(OUTPUT) / 1024:.1f} KB")


if __name__ == '__main__':
    main()
