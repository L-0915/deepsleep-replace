#!/usr/bin/env python3
"""Generate DeepSleep paper as .docx — Section 1: 背景 (complete), Sections 2-4 (placeholders)."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT_PATH = "/root/dslm/deepsleep/docs/paper.docx"


# ══════════════════════════════════════════════════════════════════════
# Helper functions
# ══════════════════════════════════════════════════════════════════════

def _set_run_font(run, font_name, size_pt, bold=False, italic=False, color=None):
    """Set font properties for a run, including East Asian font."""
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    # Set East Asian font
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        rPr = run._element.makeelement(qn('w:rPr'), {})
        run._element.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'), {})
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)


def add_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, '黑体', 18, bold=True)


def add_heading1(doc, text):
    h = doc.add_heading(level=1)
    run = h.add_run(text)
    _set_run_font(run, '黑体', 15, bold=True)


def add_heading2(doc, text):
    h = doc.add_heading(level=2)
    run = h.add_run(text)
    _set_run_font(run, '黑体', 13, bold=True)


def add_heading3(doc, text):
    h = doc.add_heading(level=3)
    run = h.add_run(text)
    _set_run_font(run, '黑体', 12, bold=True)


def add_body(doc, text):
    """Add body paragraph with first-line indent."""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run, '宋体', 12)
    return p


def add_formula(doc, text, number=None):
    """Add centered formula paragraph."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    _set_run_font(run, 'Cambria Math', 11, italic=True)
    if number:
        run2 = p.add_run(f'    ({number})')
        _set_run_font(run2, 'Cambria Math', 11)
    return p


def add_figure(doc, filename, caption):
    """Add figure placeholder with caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    run = p.add_run(f'[图: {filename}]')
    _set_run_font(run, '宋体', 10, color=RGBColor(128, 128, 128))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cap = cap.add_run(caption)
    _set_run_font(run_cap, '宋体', 10)


def add_table_caption(doc, text):
    """Add centered table caption."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, '宋体', 10, bold=True)


def add_table(doc, headers, rows):
    """Add formatted table with headers and rows."""
    table = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run_font(run, '宋体', 10, bold=True)
    # Data rows
    for r, row_data in enumerate(rows):
        for c, val in enumerate(row_data):
            cell = table.rows[r + 1].cells[c]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_run_font(run, '宋体', 10)
    return table


# ══════════════════════════════════════════════════════════════════════
# Create Document
# ══════════════════════════════════════════════════════════════════════

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

# ── Title ──────────────────────────────────────────────────────────────
add_title(doc, '基于2²全因子设计的DeepSleep轻量级')
add_title(doc, '睡眠健康语言模型DPO对齐实验研究')

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('《科学实验分析》课程大作业')
_set_run_font(run, '宋体', 12)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# Section 1: 背景
# ══════════════════════════════════════════════════════════════════════

add_heading1(doc, '1  背景')

# ── 1.1 研究背景 ──────────────────────────────────────────────────────
add_heading2(doc, '1.1  研究背景')

add_body(doc,
    '近年来，大语言模型（Large Language Model, LLM）在自然语言处理领域取得了突破性进展。'
    'Brown等[1]提出的GPT-3以1750亿参数展示了少样本学习的强大能力；OpenAI发布的GPT-4[2]'
    '进一步展现了多模态理解和复杂推理的能力；开源社区同样蓬勃发展，Meta提出的LLaMA系列[3]'
    '和阿里巴巴的Qwen系列[4]为学术研究提供了高质量的开源基座模型。这些模型在文本生成、'
    '知识问答、代码编写等通用任务上已达到令人瞩目的水平。')

add_body(doc,
    '然而，通用大语言模型在垂直领域（尤其是医疗健康）的应用仍面临诸多挑战。医疗领域具有'
    '专业知识壁垒高、术语体系复杂、安全敏感性强的特点，通用模型往往难以提供准确可靠的医学'
    '建议。为此，研究者们开展了大量医疗领域大语言模型的工作。Singhal等[5]提出Med-PaLM，'
    '基于PaLM在MultiMedQA基准上进行指令微调，首次在USMLE（美国医师执照考试）风格问题上'
    '达到"通过"水平，准确率为67.6%。其后继工作Med-PaLM 2[6]基于PaLM 2，在医学问答上'
    '达到了专家级别的性能，USMLE风格问题准确率提升至86.5%。在中文医疗领域，Zhang等[7]'
    '提出华佗GPT（HuatuoGPT），通过混合ChatGPT蒸馏数据与真实医患对话实现中文医学知识对齐；'
    'Li等[8]基于LLaMA进行医学知识微调提出ChatDoctor，显著改善了患者查询理解能力。'
    'Zhang等[9]进一步提出BiomedGPT，构建了统一的多模态生物医学基础模型。')

add_body(doc,
    '睡眠健康是医疗健康领域中一个重要且独特的研究方向。世界卫生组织数据显示，全球约30%的'
    '人口受到不同程度的睡眠障碍影响，而长期的睡眠问题与心血管疾病、糖尿病、抑郁症、阿尔茨海默病'
    '等多种慢性疾病密切相关。近年来，人工智能在睡眠健康领域的应用逐渐兴起。2025年，Bhatt等[10]'
    '提出PH-LLM（Personal Health Large Language Model），基于Gemini微调构建个人健康大语言'
    '模型，创建了针对睡眠和健身的三个评估基准。2026年，斯坦福大学医学院提出的SleepFM[11]是'
    '首个利用多导睡眠图（Polysomnography, PSG）数据预测100余种健康风险的AI模型，涵盖心血管'
    '疾病、癌症、神经退行性疾病和精神疾病。同年，JMIR发表的一项随机对照试验[12]研究表明，基于'
    'AI的睡眠改善聊天机器人具有良好的用户接受度，且与睡眠结果的改善显著相关。')

add_body(doc,
    '尽管上述研究在医疗和睡眠健康领域取得了显著进展，但仍存在以下不足：（1）现有域特定LLM的'
    '训练通常采用单一的模型架构（稠密Dense或稀疏MoE），缺乏不同架构在域特定对齐任务上的系统性'
    '对比实验；（2）对齐超参数（如DPO中的β值）的选择缺乏实验设计方法论的指导，多依赖经验调参，'
    '缺乏统计学上的严格验证；（3）轻量级模型（参数量<5亿）在资源受限场景下的对齐效果鲜有系统研究，'
    '而这对于个人健康助手等端侧部署场景至关重要。针对上述不足，本研究以睡眠健康为垂直领域，'
    '采用2²全因子实验设计方法，系统研究模型架构（MoE vs Dense）和DPO对齐强度（β参数）'
    '两个因素及其交互效应对域特定语言模型对齐效果的影响。')

# ── 1.2 模型架构 ──────────────────────────────────────────────────────
add_heading2(doc, '1.2  模型架构：稀疏混合专家（MoE）与稠密模型')
add_heading3(doc, '1.2.1  MoE架构原理')

add_body(doc,
    '稀疏混合专家（Mixture of Experts, MoE）架构的核心思想是将Transformer中的前馈网络'
    '（Feed-Forward Network, FFN）层替换为多个并行的"专家"子网络，通过门控路由机制'
    '（Gating/Router）为每个输入token选择最相关的k个专家进行处理。给定输入向量x，'
    'MoE层的输出可表示为：')

add_formula(doc, 'y = Σᵢ₌₁ᴺ G(x)ᵢ · Eᵢ(x)', '1')

add_body(doc,
    '其中N为专家总数，Eᵢ(x)为第i个专家的前馈变换，G(x)为门控函数。门控函数通常采用'
    '基于softmax的路由策略：')

add_formula(doc, 'G(x) = Softmax(Wg · x)', '2')

add_body(doc,
    '其中Wg ∈ Rᴺˣᵈ为可训练的门控权重矩阵。Top-k路由机制仅保留G(x)中得分最高的k个值，'
    '其余置零并重新归一化，实现计算资源的稀疏分配。为促进各专家的负载均衡，MoE通常引入'
    '辅助负载均衡损失（Auxiliary Load Balancing Loss）：')

add_formula(doc, 'L_aux = α · N · Σᵢ₌₁ᴺ fᵢ · Pᵢ', '3')

add_body(doc,
    '其中fᵢ为分配给第i个专家的token比例，Pᵢ为门控函数对第i个专家的平均路由概率，α为'
    '辅助损失系数。当fᵢ和Pᵢ均匀分布时L_aux最小，从而鼓励各专家均匀接收token。'
    '这种稀疏激活机制使得MoE模型能够在总参数量较大的情况下，保持每个token的计算开销较低。')

add_heading3(doc, '1.2.2  MoE的发展与应用')

add_body(doc,
    'MoE架构在大语言模型中的应用经历了快速发展。Lepikhin等[13]提出GShard，首次将MoE架构'
    '扩展至6000亿参数规模，使用top-2路由策略和辅助负载均衡损失，实现了跨数千TPU设备的'
    '自动分片训练。Fedus等[14]进一步提出Switch Transformer，将路由策略简化为top-1，'
    '成功将模型扩展至万亿参数，训练速度相比稠密模型提升了7倍。Jiang等[15]提出的Mixtral 8x7B'
    '是开源MoE大语言模型的里程碑之作，该模型总参数约467亿，每个token仅激活约129亿参数'
    '（27.6%利用率），性能匹配甚至超越了LLaMA 2 700亿参数版本和GPT-3.5。Qwen团队[4]'
    '推出的Qwen2.5系列同样包含MoE变体，在超过20万亿tokens上进行了预训练，经SFT和RLHF'
    '对齐后在多项基准上达到领先水平。')

add_heading3(doc, '1.2.3  MoE对齐的挑战')

add_body(doc,
    '尽管MoE架构在预训练和推理效率方面表现优异，但在偏好对齐（Preference Alignment）阶段'
    '面临独特挑战：')

add_body(doc,
    '（1）专家坍塌（Expert Collapse）：路由机制可能收敛到反复使用少量"优势"专家，而其他'
    '专家由于缺乏梯度信号逐渐"死亡"，导致MoE退化为较小的稠密模型。He等[16]的研究指出，'
    'MoE模型在强化学习训练中可能遭遇灾难性的训练崩溃，这与路由机制的数值不稳定性和梯度稀疏性'
    '密切相关。')

add_body(doc,
    '（2）训练不稳定性：相比稠密模型，MoE在偏好优化阶段的路由梯度更加稀疏。由于每个token'
    '仅激活top-k个专家，未被选中的专家在当前步中不接收梯度更新。这种稀疏梯度信号在偏好优化'
    '的小学习率设置下可能导致训练不稳定。')

add_body(doc,
    '（3）路由噪声：门控网络的随机性（如训练时添加的router jitter noise）在偏好优化中引入'
    '额外的方差，可能导致对齐效果的波动。')

add_body(doc,
    '相比之下，稠密（Dense）架构虽然参数效率较低（每个token需要激活全部参数），但由于所有参数'
    '在每个token上都被激活，梯度信号更加均匀，训练过程更加稳定。因此，系统对比两种架构在域特定'
    '对齐任务上的表现差异，具有重要的理论和实践价值。')

# ── 1.3 DPO对齐方法 ──────────────────────────────────────────────────
add_heading2(doc, '1.3  DPO对齐方法')
add_heading3(doc, '1.3.1  从RLHF到DPO')

add_body(doc,
    '传统的大语言模型对齐流程采用基于人类反馈的强化学习（Reinforcement Learning from Human '
    'Feedback, RLHF）[17]，其标准流程包含三个阶段：（1）监督微调（Supervised Fine-Tuning, '
    'SFT），使用高质量对话数据进行指令遵循能力训练；（2）训练奖励模型（Reward Model, RM），'
    '基于人类偏好数据学习一个评分函数rφ(y, x)；（3）使用PPO（Proximal Policy Optimization）'
    '等强化学习算法，以奖励模型的得分为信号优化语言模型策略。RLHF的优化目标为：')

add_formula(doc, 'max_πθ E_x~D,y~πθ [rφ(y, x)] - β_KL · D_KL(πθ || πref)', '4')

add_body(doc,
    '其中β_KL为KL散度惩罚系数，用于约束策略模型不偏离参考模型过远。然而，RLHF存在诸多'
    '实际问题：训练过程不稳定（RL策略可能发散）、奖励模型容易过拟合、计算成本高昂（需要同时'
    '维护策略模型、参考模型、奖励模型和价值网络四个模型）。')

add_body(doc, '在SFT阶段，训练目标为标准的自回归交叉熵损失：')

add_formula(doc, 'L_SFT(θ) = -Σₜ₌₁ᵀ log π_θ(yₜ | x, y<ₜ)', '5')

add_body(doc,
    '其中x为输入prompt，yₜ为第t个目标token，π_θ为模型在参数θ下的条件概率分布。')

add_heading3(doc, '1.3.2  DPO损失函数')

add_body(doc,
    'Rafailov等[18]提出直接偏好优化（Direct Preference Optimization, DPO），通过重新参数化'
    '将RLHF问题转化为简单的二元分类问题，绕过了显式奖励模型和强化学习的复杂流程。DPO的理论'
    '出发点是：在最优策略下，奖励函数可以被表示为策略与参考模型之间的对数概率比。给定偏好数据集'
    'D = {(xᵢ, yʷᵢ, yˡᵢ)}，其中xᵢ为输入prompt，yʷᵢ为偏好的回复（chosen/winner），'
    'yˡᵢ为不偏好的回复（rejected/loser），DPO定义隐式奖励函数为：')

add_formula(doc, 'r̂_θ(y, x) = β · log(π_θ(y|x) / π_ref(y|x))', '6')

add_body(doc,
    '其中π_θ为待优化的策略模型，π_ref为参考模型（通常为SFT后的模型快照，在DPO训练中参数冻结）。'
    '基于Bradley-Terry偏好模型，DPO的损失函数为：')

add_formula(doc,
    'L_DPO(π_θ) = -E_(x,yw,yl)~D [log σ(β · log(π_θ(yw|x)/π_ref(yw|x)) '
    '- β · log(π_θ(yl|x)/π_ref(yl|x)))]',
    '7')

add_body(doc,
    '其中σ(·)为logistic sigmoid函数σ(z) = 1/(1 + e⁻ᶻ)。该公式可简化为隐式奖励差的形式：')

add_formula(doc, 'L_DPO = -E [log σ(r̂_θ(yw, x) - r̂_θ(yl, x))]', '8')

add_body(doc,
    'DPO的核心思想是将奖励函数重新参数化为策略模型与参考模型之间的对数概率比，从而将'
    '"学习奖励函数+强化学习优化"的两步流程简化为一步直接优化。当π_θ分配给偏好回复yw的概率'
    '高于参考模型π_ref时，r̂_θ(yw, x)为正，反之为负。式(7)中的梯度为：')

add_formula(doc,
    '∇_θ L_DPO = -E [σ(r̂_l - r̂_w) · β · (∇_θ log(π_θ(yw|x)/π_θ(yl|x)))]',
    '9')

add_body(doc,
    '其中r̂_w = r̂_θ(yw, x)，r̂_l = r̂_θ(yl, x)。当模型正确区分偏好（即r̂_w >> r̂_l）时，'
    'σ(r̂_l - r̂_w)趋近于0，梯度消失，训练自然收敛。')

add_heading3(doc, '1.3.3  Beta参数的作用机制')

add_body(doc,
    'Beta参数（β）在DPO中起着关键的调节作用，其物理含义可以从多个角度理解：')

add_body(doc,
    '（1）温度缩放：β控制隐式奖励r̂_θ的温度。β值越高，模型对偏好差异（chosen与rejected之间'
    '的对数概率差）越敏感，对齐信号越强。')

add_body(doc,
    '（2）KL约束的逆温度：在理论框架中，β与策略模型偏离参考模型的KL散度惩罚有关。较小的β'
    '允许策略更大地偏离参考模型，较大的β则使优化更加保守。但在实现层面，较大的β实际上加速了'
    '训练收敛，因为sigmoid函数在更大输入值下更快饱和至1，使得损失更快接近0。')

add_body(doc,
    '（3）常用范围与推荐值：Hugging Face的系统实验[19]和Together AI的工程实践均确认'
    'β ∈ [0.1, 0.5]为有效区间，其中β = 0.1为最常用的默认值，β = 0.5代表更强的对齐信号。'
    'Rafailov等[18]在原始DPO论文中同样推荐β ∈ [0.1, 0.5]的设置范围。')

add_body(doc,
    '后续改进工作进一步拓展了DPO的理论和实践：Azar等[20]提出IPO（Identity Preference '
    'Optimization），用正则化参数τ替代β，理论框架更加稳健；Ethayarajh等[21]提出KTO'
    '（Kahneman-Tversky Optimization），仅需二值反馈（好/坏）而非成对偏好数据，降低了数据'
    '标注成本；Hong等[22]提出ORPO（Odds Ratio Preference Optimization），结合odds ratio'
    '权重，无需参考模型，进一步简化了训练流程；Wang等[23]提出beta-DPO，引入动态beta调度策略，'
    '根据训练阶段自适应调整对齐强度。')

# ── 1.4 实验目的 ──────────────────────────────────────────────────────
add_heading2(doc, '1.4  实验目的')

add_body(doc,
    '本实验以睡眠健康为垂直领域，构建轻量级MoE语言模型DeepSleep（总参数约1.99亿，每token'
    '活跃参数约6450万，词表大小7200），以Qwen2.5-0.5B（稠密模型，总参数约4.94亿）为对比'
    '基线，通过2²全因子实验设计系统研究以下问题：')

add_body(doc,
    '（1）模型架构因素（因素A）：DeepSleep稀疏MoE架构与Qwen2.5稠密架构在域特定DPO对齐效果'
    '上的差异；')

add_body(doc,
    '（2）DPO Beta参数因素（因素B）：β = 0.1（弱对齐）与β = 0.5（强对齐）对对齐效果的'
    '影响；')

add_body(doc,
    '（3）交互效应（A×B）：最优Beta值是否因模型架构而异——即模型架构与对齐强度之间是否存在'
    '统计显著的交互作用。')

add_body(doc,
    '通过量化主效应和交互效应的大小和统计显著性，为轻量级域特定语言模型的对齐超参数选择提供'
    '基于实验设计方法论的实证依据。')

# ── 1.5 实验响应变量 ──────────────────────────────────────────────────
add_heading2(doc, '1.5  实验响应变量')

add_body(doc,
    '本实验定义以下三个主要响应变量（评价指标），从不同维度衡量DPO对齐效果：')

add_table_caption(doc, '表1  实验响应变量定义')
add_table(doc,
    ['响应变量', '符号', '定义', '优化方向', '采集方式'],
    [
        ['DPO最终收敛损失', 'Y₁', 'DPO训练结束时的最终loss值', '越小越好', '训练日志自动记录'],
        ['损失下降比', 'Y₂', '(loss₅₀ − loss_final) / loss₅₀', '越大越好', '计算得出'],
        ['归一化AUC', 'Y₃', '训练loss曲线下面积的归一化值', '越小越好', '数值积分'],
    ])

add_body(doc,
    '其中，Y₁（DPO最终收敛损失）直接反映训练的拟合程度，loss越低表示策略模型与偏好数据的'
    '一致性越高。Y₂（损失下降比）衡量训练过程中loss的相对改善幅度，消除了不同模型初始loss'
    '差异的影响：')

add_formula(doc, 'Y₂ = (L₅₀ - L_final) / L₅₀', '10')

add_body(doc,
    '其中L₅₀为第50步的训练loss（代表训练初期的loss水平），L_final为训练结束时的最终loss。'
    'Y₂越接近1表示训练改善幅度越大。')

add_body(doc,
    'Y₃（归一化AUC）反映训练收敛速度，通过对训练loss曲线进行数值积分得到曲线下面积（AUC），'
    '并除以总步数进行归一化。AUC越小表示模型越快达到低loss水平：')

add_formula(doc, 'Y₃ = (1/T) · Σₜ₌₁ᵀ L(t)', '11')

add_body(doc,
    '其中L(t)为第t步的训练loss，T为总训练步数。')

add_body(doc,
    '此外，作为补充评估，还对DPO后的模型进行：（1）5个标准Benchmark评估（PubMedQA、MedQA、'
    'ARC-Easy、PIQA、OpenBookQA），使用lm-evaluation-harness框架测量acc_norm指标，采用'
    '无重复双因素方差分析（Two-Way ANOVA without replication）检验模型间差异的统计显著性；'
    '（2）30条统一测试prompt的10维度生成质量评估（专业性、安全性、人格一致性、实用性、同理心等），'
    '由DeepSeek V4模型进行自动评分，并使用单因素方差分析（One-Way ANOVA）检验模型间评分差异。')


# ══════════════════════════════════════════════════════════════════════
# Placeholder: Sections 2-4
# ══════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading1(doc, '2  实验设计')
add_body(doc, '（待续）')

doc.add_page_break()
add_heading1(doc, '3  实验结果和处理')
add_body(doc, '（待续）')

doc.add_page_break()
add_heading1(doc, '4  结果与讨论')
add_body(doc, '（待续）')

# ══════════════════════════════════════════════════════════════════════
# References
# ══════════════════════════════════════════════════════════════════════

doc.add_page_break()
add_heading1(doc, '参考文献')

refs = [
    '[1] Brown T, Mann B, Ryder N, et al. Language Models are Few-Shot Learners[C]. '
    'Advances in Neural Information Processing Systems (NeurIPS), 2020, 33: 1877-1901.',

    '[2] OpenAI. GPT-4 Technical Report[J]. arXiv preprint arXiv:2303.08774, 2023.',

    '[3] Touvron H, Lavril T, Izacard G, et al. LLaMA: Open and Efficient Foundation '
    'Language Models[J]. arXiv preprint arXiv:2302.13971, 2023.',

    '[4] Yang A, Yang B, Hui B, et al. Qwen2.5 Technical Report[J]. '
    'arXiv preprint arXiv:2412.15115, 2024.',

    '[5] Singhal K, Azizi S, Tu T, et al. Large Language Models Encode Clinical '
    'Knowledge[J]. Nature, 2023, 620: 172-179.',

    '[6] Singhal K, Tu T, Gottweis J, et al. Towards Expert-Level Medical Question '
    'Answering with Large Language Models[J]. arXiv preprint arXiv:2405.14027, 2024.',

    '[7] Zhang H, Chen D, Gao C, et al. HuatuoGPT: Towards Taming Language Models '
    'to Be a Doctor[C]. Findings of EMNLP 2023, pp. 10884-10901.',

    '[8] Li Y, Wang S, Ding H, Chen H. ChatDoctor: A Medical Chat Model Fine-Tuned '
    'on LLaMA Using Medical Domain Knowledge[J]. arXiv preprint arXiv:2303.14070, 2023.',

    '[9] Zhang K, Li J, Li D, et al. BiomedGPT: A Unified and Generalist Biomedical '
    'Generative Pre-trained Transformer for Vision, Language, and Multimodal Tasks[J]. '
    'arXiv preprint arXiv:2305.17100, 2023.',

    '[10] Bhatt M, Chou L, Gao J, et al. A Personal Health Large Language Model '
    'for Sleep and Fitness[J]. Nature Medicine, 2025.',

    '[11] SleepFM: Multi-Modal Sleep Foundation Model for Predicting 100+ Health '
    'Risks[R]. Stanford Medicine, 2026.',

    '[12] AI Chatbot-Delivered Intervention for Sleep Improvement: Randomized '
    'Controlled Trial[J]. Journal of Medical Internet Research (JMIR), 2026.',

    '[13] Lepikhin D, Lee H, Xu Y, et al. GShard: Scaling Giant Models with '
    'Conditional Computation and Automatic Sharding[C]. ICLR 2021.',

    '[14] Fedus W, Zoph B, Shazeer N. Switch Transformers: Scaling to Trillion '
    'Parameter Models with Simple and Efficient Sparsity[J]. Journal of Machine '
    'Learning Research, 2022, 23(120): 1-39.',

    '[15] Jiang A Q, Sablayrolles A, Roux A, et al. Mixtral of Experts[J]. '
    'arXiv preprint arXiv:2401.04088, 2024.',

    '[16] He J, Qiu C, et al. MoE Training Instability and Its Remedies in '
    'Reinforcement Learning[J]. arXiv preprint arXiv:2510.11370, 2025.',

    '[17] Ouyang L, Wu J, Jiang X, et al. Training language models to follow '
    'instructions with human feedback[C]. Advances in Neural Information Processing '
    'Systems (NeurIPS), 2022, 35: 27730-27744.',

    '[18] Rafailov R, Sharma A, Mitchell E, et al. Direct Preference Optimization: '
    'Your Model is Secretly a Reward Model[C]. Advances in Neural Information '
    'Processing Systems (NeurIPS), 2023.',

    '[19] Hugging Face. Preference Tuning LLMs with Direct Preference Optimization '
    'Methods[EB/OL]. https://huggingface.co/blog/pref-tuning, 2024.',

    '[20] Azar M G, Guo Z D, Piot B, et al. A General Theoretical Paradigm to '
    'Understand Learning from Human Preferences[C]. AISTATS 2024.',

    '[21] Ethayarajh K, Xu W, Zhang N, et al. KTO: Model Alignment as Prospect '
    'Theoretic Optimization[J]. arXiv preprint arXiv:2402.01306, 2024.',

    '[22] Hong J, Lee H, Lee J, et al. ORPO: Monolithic Preference Optimization '
    'without Reference Model[J]. arXiv preprint arXiv:2403.07691, 2024.',

    '[23] Wang Z, et al. Beta-DPO: Dynamic Beta Scheduling for Direct Preference '
    'Optimization[J]. arXiv preprint arXiv:2407.08639, 2024.',
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(ref)
    _set_run_font(run, '宋体', 10)

# ══════════════════════════════════════════════════════════════════════
# Save
# ══════════════════════════════════════════════════════════════════════

os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)
print(f"Done: {OUT_PATH}")
